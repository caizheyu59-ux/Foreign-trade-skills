#!/usr/bin/env python3
"""
报价单翻译器 - 读取现有报价单，仅翻译文字内容，保留原始样式和布局
支持格式：Excel (.xlsx/.xls)、HTML、Word (.docx)
支持语言：zh, en, es, hi, ar, fr, ru, pt, de, ja, ko

v2.0.0 - 直接翻译模式，无需外部 API
"""

import sys
import os
import json
from pathlib import Path
import argparse

# 多语言配置
SUPPORTED_LANGUAGES = ['zh', 'en', 'es', 'hi', 'ar', 'fr', 'ru', 'pt', 'de', 'ja', 'ko']

LANGUAGE_NAMES = {
    'zh': '中文', 'en': 'English', 'es': 'Español', 'hi': 'हिन्दी',
    'ar': 'العربية', 'fr': 'Français', 'ru': 'Русский', 'pt': 'Português',
    'de': 'Deutsch', 'ja': '日本語', 'ko': '한국어'
}

# 通用翻译词库（可扩展）
BASE_TRANSLATIONS = {
    # 英语 → 其他语言
    'en': {
        'Company': '公司', 'Date': '日期', 'Product': '产品', 'Price': '价格',
        'Total': '总计', 'Quotation': '报价单', 'Contact': '联系',
        'Charms': '吊坠', 'Stainless steel Charms': '不锈钢吊坠',
        'Contact us for more Collection': '联系我们获取更多系列',
        '1688 Link': '1688 链接',
    },
    # 英语 → 日语
    'ja': {
        'Company': '会社', 'Date': '日付', 'Product': '製品', 'Price': '価格',
        'Total': '合計', 'Quotation': '見積書', 'Contact': '連絡先',
        'Charms': 'チャーム', 'Stainless steel Charms': 'ステンレススチールチャーム',
        'Contact us for more Collection': 'その他のコレクションについてはお問い合わせください',
        '1688 Link': '1688 リンク',
    },
    # 英语 → 西班牙语
    'es': {
        'Company': 'Empresa', 'Date': 'Fecha', 'Product': 'Producto', 'Price': 'Precio',
        'Total': 'Total', 'Quotation': 'Cotización', 'Contact': 'Contacto',
        'Charms': 'Dijes', 'Stainless steel Charms': 'Dijes de acero inoxidable',
        'Contact us for more Collection': 'Contáctenos para más colección',
        '1688 Link': 'Enlace 1688',
    },
    # 英语 → 法语
    'fr': {
        'Company': 'Société', 'Date': 'Date', 'Product': 'Produit', 'Price': 'Prix',
        'Total': 'Total', 'Quotation': 'Devis', 'Contact': 'Contact',
        'Charms': 'Breloques', 'Stainless steel Charms': 'Breloques en acier inoxydable',
        'Contact us for more Collection': 'Contactez-nous pour plus de collection',
        '1688 Link': 'Lien 1688',
    },
    # 英语 → 德语
    'de': {
        'Company': 'Firma', 'Date': 'Datum', 'Product': 'Produkt', 'Price': 'Preis',
        'Total': 'Gesamt', 'Quotation': 'Angebot', 'Contact': 'Kontakt',
        'Charms': 'Anhänger', 'Stainless steel Charms': 'Anhänger aus Edelstahl',
        'Contact us for more Collection': 'Kontaktieren Sie uns für weitere Kollektionen',
        '1688 Link': '1688 Link',
    },
    # 英语 → 韩语
    'ko': {
        'Company': '회사', 'Date': '날짜', 'Product': '제품', 'Price': '가격',
        'Total': '합계', 'Quotation': '견적서', 'Contact': '연락처',
        'Charms': '참', 'Stainless steel Charms': '스테인리스 스틸 참',
        'Contact us for more Collection': '더 많은 컬렉션은 문의하세요',
        '1688 Link': '1688 링크',
    },
    # 英语 → 俄语
    'ru': {
        'Company': 'Компания', 'Date': 'Дата', 'Product': 'Товар', 'Price': 'Цена',
        'Total': 'Итого', 'Quotation': 'Коммерческое предложение', 'Contact': 'Контакт',
        'Charms': 'Подвески', 'Stainless steel Charms': 'Подвески из нержавеющей стали',
        'Contact us for more Collection': 'Свяжитесь с нами для получения дополнительной коллекции',
        '1688 Link': 'Ссылка 1688',
    },
    # 英语 → 阿拉伯语
    'ar': {
        'Company': 'شركة', 'Date': 'تاريخ', 'Product': 'منتج', 'Price': 'سعر',
        'Total': 'المجموع', 'Quotation': 'عرض أسعار', 'Contact': 'اتصل بنا',
        'Charms': 'الحلي', 'Stainless steel Charms': 'الحلي الفولاذ المقاوم للصدأ',
        'Contact us for more Collection': 'اتصل بنا للحصول على المزيد من المجموعة',
        '1688 Link': 'رابط 1688',
    },
}

# 长文本翻译模板（按行业/场景）
LONG_TEXT_TRANSLATIONS = {
    # 珠宝行业描述
    'As a jewelry, a pendant not only has a decorative role, but also carries a rich moral and symbolic meaning. Here are some pendants and charms.\n\nWith our rich experience, OEM & ODM services are also welcome here, send us your design or idea, we will make your idea come into reality.': {
        'zh': '作为珠宝，吊坠不仅具有装饰作用，还承载着丰富的寓意和象征意义。以下是一些吊坠和挂饰。\n\n凭借我们丰富的经验，我们也欢迎 OEM 和 ODM 服务，发送您的设计或想法给我们，我们将使您的想法成为现实。',
        'ja': 'ジュエリーとして、ペンダントは装飾的な役割だけでなく、豊かな道徳的・象徴的な意味を持っています。こちらはいくつかのペンダントとチャームです。\n\n私たちの豊富な経験により、OEM および ODM サービスも歓迎いたします。デザインやアイデアをお送りください。あなたのアイデアを現実にします。',
        'es': 'Como joyería, un colgante no solo tiene un papel decorativo, sino que también lleva un rico significado moral y simbólico. Aquí hay algunos colgantes y dijes.\n\nCon nuestra rica experiencia, los servicios de OEM y ODM también son bienvenidos aquí, envíenos su diseño o idea, haremos que su idea se haga realidad.',
        'fr': 'En tant que bijou, un pendentif a non seulement un rôle décoratif, mais porte également une riche signification morale et symbolique. Voici quelques pendentifs et breloques.\n\nAvec notre riche expérience, les services OEM et ODM sont également les bienvenus ici, envoyez-nous votre conception ou idée, nous ferons de votre idée une réalité.',
        'de': 'Als Schmuckstück hat ein Anhänger nicht nur eine dekorative Rolle, sondern trägt auch eine reiche moralische und symbolische Bedeutung. Hier sind einige Anhänger und Charms.\n\nMit unserer reichen Erfahrung sind OEM- und ODM-Dienstleistungen auch hier willkommen, senden Sie uns Ihr Design oder Ihre Idee, wir werden Ihre Idee verwirklichen.',
        'ko': '주얼리로서 펜던트는 장식적인 역할뿐만 아니라 풍부한 도덕적, 상징적 의미를 담고 있습니다. 여기 몇 가지 펜던트와 참이 있습니다.\n\n우리의 풍부한 경험을 바탕으로 OEM 및 ODM 서비스도 환영합니다. 디자인이나 아이디어를 보내주시면 아이디어를 현실로 만들어 드립니다.',
        'ru': 'Как украшение, кулон не только играет декоративную роль, но и несет богатое моральное и символическое значение. Вот некоторые кулоны и подвески.\n\nБлагодаря нашему богатому опыту, услуги OEM и ODM также приветствуются здесь, отправьте нам свой дизайн или идею, мы воплотим вашу идею в реальность.',
        'ar': 'كمجوهرات، لا يلعب القلادة دورًا زخرفيًا فحسب، بل يحمل أيضًا معنى أخلاقي ورمزي غني. إليك بعض القلائد والحلي.\n\nمع خبرتنا الغنية، خدمات OEM و ODM مرحب بها أيضًا هنا، أرسل لنا تصميمك أو فكرتك، سنجعل فكرتك تتحقق.',
    },
}


def load_custom_translations(custom_file):
    """加载自定义翻译词库"""
    if not custom_file or not Path(custom_file).exists():
        return {}
    try:
        with open(custom_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"[WARN] Failed to load custom translations: {e}")
        return {}


def translate_text(text, target_lang, custom_translations=None):
    """翻译单段文本"""
    if not text or not isinstance(text, str) or not text.strip():
        return text
    
    # 合并词库
    translations = BASE_TRANSLATIONS.get(target_lang, {})
    if custom_translations:
        custom_target = custom_translations.get(target_lang, {})
        translations = {**translations, **custom_target}
    
    # 先匹配长文本
    for src, trans_dict in LONG_TEXT_TRANSLATIONS.items():
        if text.strip() == src or text.strip() in src:
            return trans_dict.get(target_lang, text)
    
    # 精确匹配
    if text in translations:
        return translations[text]
    
    # 部分匹配（忽略大小写）
    text_lower = text.lower()
    for src, dst in translations.items():
        if text_lower == src.lower():
            return dst
    
    # 无匹配时返回原文
    return text


def should_skip_translation(text):
    """判断是否应该跳过翻译"""
    if not text or not isinstance(text, str):
        return True
    
    text = text.strip()
    if len(text) < 3:
        return True
    
    # 跳过全大写的文本（通常是产品编号、缩写等）
    if text.upper() == text:
        return True
    
    # 跳过网址、邮箱
    if '@' in text or 'www.' in text or 'http' in text or 'https' in text:
        return True
    
    # 跳过纯数字
    if text.replace('.', '').replace(',', '').replace('-', '').isdigit():
        return True
    
    return False


def translate_excel(excel_path, target_lang, custom_translations=None):
    """翻译 Excel 文件，保留所有样式、公式、格式"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        print("[ERROR] openpyxl not installed. Install: pip install openpyxl")
        sys.exit(1)
    
    # 加载工作簿，保留样式
    wb = load_workbook(excel_path, keep_vba=True)
    print(f"     Sheets found: {wb.sheetnames}")
    
    count = 0
    # 遍历所有工作表
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        print(f"     Translating sheet: {sheet_name}")
        
        # 遍历所有单元格
        for row in sheet.iter_rows():
            for cell in row:
                if not should_skip_translation(cell.value):
                    original = cell.value
                    translated = translate_text(original, target_lang, custom_translations)
                    if translated != original:
                        # 仅替换文本，保留样式、公式、格式
                        cell.value = translated
                        count += 1
    
    # 保存翻译后的文件
    output_path = str(excel_path).replace('.xlsx', f'_{target_lang}.xlsx')
    if excel_path.suffix.lower() == '.xls':
        output_path = str(excel_path).replace('.xls', f'_{target_lang}.xls')
    elif excel_path.suffix.lower() == '.xlsm':
        output_path = str(excel_path).replace('.xlsm', f'_{target_lang}.xlsm')
    
    wb.save(output_path)
    return output_path, count


def translate_html(html_path, target_lang, custom_translations=None):
    """翻译 HTML 文件，保留所有样式"""
    try:
        from bs4 import BeautifulSoup, NavigableString
    except ImportError:
        print("[ERROR] beautifulsoup4 not installed. Install: pip install beautifulsoup4 lxml")
        sys.exit(1)
    
    with open(html_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    soup = BeautifulSoup(content, 'lxml')
    
    # 阿拉伯语添加 RTL 支持
    if target_lang == 'ar' and soup.html:
        soup.html['dir'] = 'rtl'
        soup.html['lang'] = target_lang
    
    count = 0
    
    # 递归翻译所有文本节点
    def translate_node(node):
        nonlocal count
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text and not should_skip_translation(text):
                translated = translate_text(text, target_lang, custom_translations)
                if translated != text:
                    node.replace_with(translated)
                    count += 1
        else:
            # 跳过 script 和 style 标签
            if node.name not in ['script', 'style']:
                for child in node.children:
                    translate_node(child)
    
    # 从 body 开始翻译（如果有 body）
    if soup.body:
        translate_node(soup.body)
    else:
        translate_node(soup)
    
    # 保存
    output_path = str(html_path).replace('.html', f'_{target_lang}.html')
    if html_path.suffix.lower() == '.htm':
        output_path = str(html_path).replace('.htm', f'_{target_lang}.htm')
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(str(soup))
    
    return output_path, count


def translate_word(docx_path, target_lang, custom_translations=None):
    """翻译 Word 文档，保留所有样式"""
    try:
        from docx import Document
    except ImportError:
        print("[ERROR] python-docx not installed. Install: pip install python-docx")
        sys.exit(1)
    
    doc = Document(docx_path)
    count = 0
    
    # 翻译所有段落
    for para in doc.paragraphs:
        if para.text.strip() and not should_skip_translation(para.text):
            translated = translate_text(para.text, target_lang, custom_translations)
            if translated != para.text:
                runs = para.runs
                if runs:
                    for run in runs[1:]:
                        run.text = ''
                    runs[0].text = translated
                else:
                    para.clear()
                    para.add_run(translated)
                count += 1
    
    # 翻译表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip() and not should_skip_translation(para.text):
                        translated = translate_text(para.text, target_lang, custom_translations)
                        if translated != para.text:
                            runs = para.runs
                            if runs:
                                for run in runs[1:]:
                                    run.text = ''
                                runs[0].text = translated
                            else:
                                para.clear()
                                para.add_run(translated)
                            count += 1
    
    # 保存
    output_path = str(docx_path).replace('.docx', f'_{target_lang}.docx')
    doc.save(output_path)
    return output_path, count


def main():
    parser = argparse.ArgumentParser(description='报价单翻译器 - 保留原样式，仅翻译文字 (v2.0.0)')
    parser.add_argument('input', help='输入文件路径（Excel/HTML/Word）')
    parser.add_argument('-l', '--lang', default='en', choices=SUPPORTED_LANGUAGES, help='目标语言')
    parser.add_argument('-o', '--output', help='输出文件路径')
    parser.add_argument('-c', '--custom', help='自定义翻译词库 JSON 文件')
    parser.add_argument('-v', '--verbose', action='store_true', help='显示详细翻译过程')
    args = parser.parse_args()
    
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"[ERROR] File not found: {input_path}")
        sys.exit(1)
    
    # 加载自定义词库
    custom_translations = load_custom_translations(args.custom)
    if custom_translations and args.verbose:
        print(f"[INFO] Loaded custom translations from {args.custom}")
    
    print(f"[START] Translating {input_path.name} to {LANGUAGE_NAMES.get(args.lang, args.lang)}")
    
    # 根据文件类型选择翻译方式
    suffix = input_path.suffix.lower()
    
    if suffix in ['.xlsx', '.xlsm', '.xls']:
        output_path, count = translate_excel(input_path, args.lang, custom_translations)
    elif suffix in ['.html', '.htm']:
        output_path, count = translate_html(input_path, args.lang, custom_translations)
    elif suffix == '.docx':
        output_path, count = translate_word(input_path, args.lang, custom_translations)
    else:
        print(f"[ERROR] Unsupported file format: {suffix}")
        print("        Supported formats: .xlsx, .xlsm, .xls, .html, .htm, .docx")
        sys.exit(1)
    
    # 使用用户指定的输出路径
    if args.output:
        output_path = args.output
    
    print(f"[OK] Translation complete: {output_path}")
    print(f"     Target language: {LANGUAGE_NAMES.get(args.lang, args.lang)}")
    print(f"     Translated {count} cells/elements")
    
    return output_path


if __name__ == '__main__':
    main()
